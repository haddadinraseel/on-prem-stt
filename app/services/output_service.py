from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.core.constants import JOB_METADATA_NAME, TRANSCRIPT_DOCX_NAME, TRANSCRIPT_TEXT_NAME
from app.models.job import JobRecord
from app.utils.formatting import format_timestamp


class OutputService:
    def write_outputs(self, job: JobRecord, output_dir: Path) -> tuple[Path, Path, Path]:
        output_dir.mkdir(parents=True, exist_ok=True)
        text_path = output_dir / TRANSCRIPT_TEXT_NAME
        docx_path = output_dir / TRANSCRIPT_DOCX_NAME
        metadata_path = output_dir / JOB_METADATA_NAME

        lines = []
        for segment in job.segments:
            label = segment.speaker or "Timestamp"
            lines.append(
                f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}] {label}: {segment.text}"
            )

        transcript_text = "\n".join(lines)
        text_output = transcript_text
        text_path.write_text(text_output, encoding="utf-8")

        document = Document()
        document.add_heading("Transcription Result", level=1)
        document.add_paragraph(f"Job ID: {job.job_id}")
        document.add_paragraph(f"Model: {job.model_name}")
        document.add_paragraph(f"Device: {job.device or 'unknown'}")
        document.add_paragraph(f"Diarization: {job.diarization_status}")
        document.add_heading("Transcript", level=2)
        for line in lines:
            document.add_paragraph(line)
        document.save(docx_path)

        metadata_path.write_text(
            json.dumps(job.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return text_path, docx_path, metadata_path


output_service = OutputService()
